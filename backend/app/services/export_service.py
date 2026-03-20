import os
import tempfile
import re
from html import unescape
from docx import Document
import markdown
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

ORDERED_LIST_RE = re.compile(r"^\d+\.\s+")
ATX_HEADER_RE = re.compile(r"^(#{1,6})\s+(.*)$")
TABLE_DIVIDER_RE = re.compile(r"^\s*\|?[\s:-]+\|[\s|:-]*$")
MD_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
MD_IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")

class ExportService:
    @staticmethod
    def _to_plain_text(content: str) -> str:
        text = content.replace("\r\n", "\n")
        text = re.sub(r"```[\s\S]*?```", lambda m: m.group(0).strip("`"), text)
        text = MD_IMAGE_RE.sub(lambda m: m.group(1) or m.group(2), text)
        text = MD_LINK_RE.sub(r"\1 (\2)", text)
        text = re.sub(r"^\s{0,3}#{1,6}\s*", "", text, flags=re.MULTILINE)
        text = re.sub(r"^\s{0,3}>\s?", "", text, flags=re.MULTILINE)
        text = re.sub(r"^\s*[-*+]\s+", "• ", text, flags=re.MULTILINE)
        text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"(?<!\*)\*\*(.+?)\*\*(?!\*)", r"\1", text)
        text = re.sub(r"(?<!\*)\*(.+?)\*(?!\*)", r"\1", text)
        text = re.sub(r"__(.+?)__", r"\1", text)
        text = re.sub(r"_(.+?)_", r"\1", text)
        text = re.sub(r"`([^`]+)`", r"\1", text)
        text = re.sub(r"~~(.+?)~~", r"\1", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def _append_inline_runs(paragraph, text: str) -> None:
        if not text:
            paragraph.add_run("")
            return
        cursor = 0
        token_re = re.compile(r"(\*\*.+?\*\*|`.+?`|\*.+?\*)")
        for match in token_re.finditer(text):
            start, end = match.span()
            if start > cursor:
                paragraph.add_run(unescape(text[cursor:start]))
            token = match.group(0)
            if token.startswith("**") and token.endswith("**"):
                run = paragraph.add_run(unescape(token[2:-2]))
                run.bold = True
            elif token.startswith("`") and token.endswith("`"):
                run = paragraph.add_run(unescape(token[1:-1]))
                run.font.name = "Menlo"
                run.font.size = Pt(10)
            elif token.startswith("*") and token.endswith("*"):
                run = paragraph.add_run(unescape(token[1:-1]))
                run.italic = True
            cursor = end
        if cursor < len(text):
            paragraph.add_run(unescape(text[cursor:]))

    @staticmethod
    def export_to_txt(content: str) -> str:
        fd, path = tempfile.mkstemp(suffix=".txt")
        with os.fdopen(fd, 'w', encoding='utf-8') as f:
            f.write(ExportService._to_plain_text(content))
        return path

    @staticmethod
    def export_to_docx(content: str) -> str:
        doc = Document()
        lines = content.replace("\r\n", "\n").split("\n")
        in_code_block = False
        code_buffer: list[str] = []
        for raw_line in lines:
            line = raw_line.rstrip()
            stripped = line.strip()

            if stripped.startswith("```"):
                if in_code_block:
                    code_text = "\n".join(code_buffer).strip()
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = "Menlo"
                    run.font.size = Pt(10)
                    code_buffer = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue

            if in_code_block:
                code_buffer.append(line)
                continue

            if not stripped:
                doc.add_paragraph("")
                continue

            header_match = ATX_HEADER_RE.match(stripped)
            if header_match:
                level = min(len(header_match.group(1)), 4)
                text = header_match.group(2).strip()
                p = doc.add_heading(level=level)
                ExportService._append_inline_runs(p, text)
                continue

            if TABLE_DIVIDER_RE.match(stripped):
                continue

            if stripped.startswith("|") and stripped.endswith("|"):
                table_row = [cell.strip() for cell in stripped.strip("|").split("|")]
                p = doc.add_paragraph(" | ".join(table_row))
                p.paragraph_format.left_indent = Pt(18)
                continue

            if stripped.startswith(">"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(18)
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                ExportService._append_inline_runs(p, stripped.lstrip(">").strip())
                continue

            if stripped.startswith(("- ", "* ", "+ ")):
                p = doc.add_paragraph(style="List Bullet")
                ExportService._append_inline_runs(p, stripped[2:].strip())
                continue

            if ORDERED_LIST_RE.match(stripped):
                item_text = ORDERED_LIST_RE.sub("", stripped, count=1).strip()
                p = doc.add_paragraph(style="List Number")
                ExportService._append_inline_runs(p, item_text)
                continue

            p = doc.add_paragraph()
            ExportService._append_inline_runs(p, stripped)
        
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc.save(path)
        return path

    @staticmethod
    def export_to_pdf(content: str) -> str:
        html_content = markdown.markdown(content, extensions=['fenced_code', 'tables'])
        styled_html = f"""
        <html>
        <head>
            <style>
                body {{ font-family: sans-serif; line-height: 1.6; padding: 2em; color: #333; }}
                pre {{ background: #f4f4f4; padding: 1em; border-radius: 4px; white-space: pre-wrap; }}
                code {{ font-family: monospace; background: #f4f4f4; padding: 0.2em 0.4em; border-radius: 3px; }}
                h1, h2, h3, h4 {{ color: #111; margin-top: 1.5em; }}
                blockquote {{ border-left: 4px solid #ccc; margin-left: 0; padding-left: 1em; color: #666; }}
                table {{ border-collapse: collapse; width: 100%; margin-bottom: 1em; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background-color: #f4f4f4; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        fd, path = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)
        
        try:
            import sys
            # Workaround for macOS Apple Silicon Homebrew to find pango/cairo
            if sys.platform == "darwin":
                homebrew_lib = "/opt/homebrew/lib"
                if os.path.exists(homebrew_lib) and "DYLD_FALLBACK_LIBRARY_PATH" not in os.environ:
                    os.environ["DYLD_FALLBACK_LIBRARY_PATH"] = homebrew_lib
                    
            from weasyprint import HTML
            HTML(string=styled_html).write_pdf(path)
        except Exception as e:
            # Fallback or raise a descriptive error
            raise RuntimeError(f"PDF generation requires WeasyPrint system dependencies (e.g. pango, cairo). Error: {str(e)}")
            
        return path
