import os

from docx import Document

from app.services.export_service import ExportService


def test_export_to_txt_converts_markdown_to_plain_text():
    content = "# Title\n\n- Item one\n- Item two\n\nThis is **bold** with `code` and [link](https://example.com)."
    path = ExportService.export_to_txt(content)
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = f.read()
        assert "# Title" not in data
        assert "**bold**" not in data
        assert "`code`" not in data
        assert "Title" in data
        assert "Item one" in data
        assert "bold" in data
        assert "code" in data
    finally:
        if os.path.exists(path):
            os.remove(path)


def test_export_to_docx_maps_headings_and_lists():
    content = "# Main Title\n\n## Section A\n\n- Bullet one\n- Bullet two\n\n1. First\n2. Second\n\nNormal paragraph."
    path = ExportService.export_to_docx(content)
    try:
        doc = Document(path)
        paragraph_texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "Main Title" in paragraph_texts
        assert "Section A" in paragraph_texts
        assert "Bullet one" in paragraph_texts
        assert "First" in paragraph_texts
        assert "Normal paragraph." in paragraph_texts

        style_map = {p.text: p.style.name for p in doc.paragraphs if p.text.strip()}
        assert style_map["Main Title"].startswith("Heading")
        assert style_map["Section A"].startswith("Heading")
        assert style_map["Bullet one"] == "List Bullet"
        assert style_map["First"] == "List Number"
    finally:
        if os.path.exists(path):
            os.remove(path)
